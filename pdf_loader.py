"""PDF/DOCX loader with chunking for CV Ranking Agent."""

import os
import re
import uuid
import logging
from datetime import date, timezone, datetime
from pathlib import Path
from typing import List, Optional

from pypdf import PdfReader
from docx import Document

from models import CV
from ocr_processor import is_scanned_pdf, ocr_pdf, correct_ocr_errors, OCR_AVAILABLE

logger = logging.getLogger(__name__)

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100


# ── PDF extraction ────────────────────────────────────────────────────────────

def _extract_text_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n".join(pages)

    if is_scanned_pdf(text):
        logger.info(f"Detected scanned PDF: {path}")
        if OCR_AVAILABLE:
            ocr_text = ocr_pdf(path)
            if ocr_text:
                text = correct_ocr_errors(ocr_text)
                logger.info(f"OCR extracted {len(text)} characters")
            else:
                logger.warning("OCR failed, returning empty text")
        else:
            logger.warning("OCR libraries not installed. Install: pip install -r requirements-ocr.txt")

    return text


# ── DOCX extraction ───────────────────────────────────────────────────────────

def _iter_block_items(parent):
    """Yield paragraphs and tables in document order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for child in parent.element.body:
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def _extract_table_text(table) -> list:
    """Recursively extract text from a table including nested tables."""
    from docx.table import Table
    parts = []
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for nested in cell.tables:
                parts.extend(_extract_table_text(nested))
    return parts


def _extract_text_from_shapes(doc) -> list:
    """Extract text from drawing/shape text boxes (common in Naukri templates)."""
    WPS_TXBX = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'
    MC_ALTERNATE = '{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent'
    W_P = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    parts = []
    try:
        for shape in doc.element.body.iter():
            if shape.tag in (WPS_TXBX, MC_ALTERNATE):
                for p in shape.iter(W_P):
                    text = ''.join(r.text for r in p.iter(W_T) if r.text)
                    if text.strip():
                        parts.append(text)
    except Exception as e:
        logger.debug(f"Shape text extraction skipped: {e}")
    return parts


def _extract_text_docx(path: str) -> str:
    doc = Document(path)
    parts = []

    try:
        for block in _iter_block_items(doc):
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            if isinstance(block, Paragraph):
                if block.text.strip():
                    parts.append(block.text)
            elif isinstance(block, Table):
                parts.extend(_extract_table_text(block))
    except Exception as e:
        logger.debug(f"Block extraction partial failure: {e}")

    parts.extend(_extract_text_from_shapes(doc))

    try:
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf:
                    for p in hf.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)
    except Exception as e:
        logger.debug(f"Header/footer extraction skipped: {e}")

    seen, unique = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            unique.append(p)

    return "\n".join(unique)


def _extract_text(path: str) -> str:
    real_path = os.path.realpath(path)
    ext = Path(real_path).suffix.lower()
    if ext == ".pdf":
        return _extract_text_pdf(real_path)
    elif ext in (".docx", ".doc"):
        return _extract_text_docx(real_path)
    raise ValueError(f"Unsupported file type: {ext}")


# ── Name extraction ───────────────────────────────────────────────────────────

def extract_candidate_name(text: str) -> Optional[str]:
    """Extract the candidate's actual name from CV text.

    Strategy (in order of confidence):
    1. Explicit 'Name:' / 'Full Name:' label in first 30 lines
    2. First line in top 10 that looks like a personal name
    Returns None so caller can fall back to filename.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Proper name: 2-4 words, each Title-cased, only letters and spaces, no digits
    # Excludes single words, all-caps lines (company names), and lines with punctuation
    name_re = re.compile(r'^[A-Z][a-z]+(?:\s[A-Z][a-z]+){1,3}$')
    
    # Single word name pattern (for cases like "Sohebshaikh" - at least 6 chars, Title case)
    single_name_re = re.compile(r'^[A-Z][a-z]{5,}$')

    # Common non-name words that appear at top of resumes
    _EXCLUDE = {'resume', 'curriculum', 'vitae', 'profile', 'summary', 'objective',
                'contact', 'address', 'email', 'phone', 'mobile', 'linkedin',
                'engineer', 'manager', 'developer', 'analyst', 'consultant',
                'senior', 'junior', 'lead', 'head', 'director', 'officer',
                'transition', 'experience', 'education', 'skills', 'projects',
                'professional', 'personal', 'career', 'employment', 'work',
                'qualifications', 'certifications', 'achievements', 'about'}

    label_re = re.compile(r'^(?:full\s*)?name\s*[:\-]\s*(.+)$', re.IGNORECASE)

    # Strategy 1: explicit label
    for line in lines[:30]:
        m = label_re.match(line)
        if m:
            candidate = m.group(1).strip()
            # Remove job titles from the end
            candidate = re.sub(r'\s+(Transition|Senior|Junior|Lead|Manager|Analyst|Engineer|Developer|Consultant).*$', '', candidate, flags=re.IGNORECASE)
            if name_re.match(candidate) and candidate.lower().split()[0] not in _EXCLUDE:
                return candidate

    # Strategy 2: first line in top 10 that looks like a personal name (skip pure numbers)
    for line in lines[:10]:
        # Skip lines that are just numbers or start with numbers
        if re.match(r'^\d+$', line) or re.match(r'^\d+\s', line):
            continue
        
        # Skip common section headers (case-insensitive check)
        line_lower = line.lower()
        if any(excl in line_lower for excl in ['summary', 'profile', 'objective', 'experience', 'education', 'skills']):
            continue
        
        # Remove job titles/designations from the line
        cleaned = re.sub(r'\s+(Transition|Senior|Junior|Lead|Manager|Analyst|Engineer|Developer|Consultant).*$', '', line, flags=re.IGNORECASE).strip()
        
        # Check for proper multi-word name
        if name_re.match(cleaned) and cleaned.lower().split()[0] not in _EXCLUDE:
            return cleaned
        
        # Check for single-word name (fallback for concatenated names)
        if single_name_re.match(cleaned) and cleaned.lower() not in _EXCLUDE:
            return cleaned

    return None


# ── Experience extraction ─────────────────────────────────────────────────────

def extract_experience_years(text: str) -> Optional[float]:
    """Parse total years of professional work experience from raw CV text.
    Only counts actual employment. Excludes internships, education, courses.
    Validates results to prevent absurd values (e.g. thousands of years).
    """
    today = datetime.now(timezone.utc).date()
    MAX_REASONABLE_YEARS = 50  # sanity cap

    MONTHS = {'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
               'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}

    def _parse_date(s: str) -> Optional[date]:
        s = s.strip().lower()
        if s in ('present', 'current', 'till date', 'till now', 'ongoing'):
            return today
        m = re.match(r'([a-z]+)[\s,]+(\d{4})', s)
        if m:
            mon = MONTHS.get(m.group(1)[:3])
            yr = int(m.group(2))
            if mon and 1980 <= yr <= today.year + 1:
                return date(yr, mon, 1)
        m = re.match(r'^(\d{4})$', s)
        if m:
            yr = int(m.group(1))
            if 1980 <= yr <= today.year + 1:
                return date(yr, 1, 1)
        m = re.match(r'(\d{1,2})[/-](\d{4})', s)
        if m:
            yr = int(m.group(2))
            mon = int(m.group(1))
            if 1980 <= yr <= today.year + 1 and 1 <= mon <= 12:
                return date(yr, mon, 1)
        return None

    # Strategy 1: Naukri style header — "6y 4m" or "6 Years 4 Months"
    # Must appear near top of text (first 500 chars) to avoid false matches
    header_text = text[:500]
    m = re.search(r'\b(\d{1,2})\s*[Yy](?:ears?|rs?)?[\s,]*(\d{1,2})?\s*[Mm](?:onths?)?\b', header_text)
    if m:
        years = int(m.group(1))
        months = int(m.group(2)) if m.group(2) else 0
        result = round(years + months / 12, 1)
        if result <= MAX_REASONABLE_YEARS:
            return result

    # Strategy 2: Explicit statement — "X years of experience" / "X+ years"
    m = re.search(
        r'\b(\d{1,2}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)\b',
        text, re.IGNORECASE
    )
    if m:
        val = float(m.group(1))
        if val <= MAX_REASONABLE_YEARS:
            return val

    # Strategy 3: Sum professional employment date ranges only
    lines = text.splitlines()

    work_section_re = re.compile(
        r'^(professional\s+experience|work\s+experience|employment|experience|career\s+history|work\s+history)',
        re.IGNORECASE
    )
    non_work_section_re = re.compile(
        r'^(education|academic|internship|training|certification|course|project|publication|achievement|extra|volunteer)',
        re.IGNORECASE
    )

    in_work_section = False
    work_lines = []
    for line in lines:
        stripped = line.strip()
        if work_section_re.match(stripped):
            in_work_section = True
            continue
        if in_work_section and non_work_section_re.match(stripped):
            in_work_section = False
        if in_work_section:
            work_lines.append(stripped)

    if not work_lines:
        skip_re = re.compile(
            r'intern|trainee|apprentice|bachelor|master|b\.?tech|m\.?tech|b\.?e\b|m\.?e\b|pgdm|mba|phd|'
            r'university|college|institute|school',
            re.IGNORECASE
        )
        work_lines = [l.strip() for l in lines if not skip_re.search(l)]

    work_text = '\n'.join(work_lines)

    date_range_pattern = re.compile(
        r'([A-Za-z]+[\s,]+\d{4}|\d{4}|\d{1,2}[/-]\d{4})'
        r'\s*[-\u2013\u2014to]+\s*'
        r'([A-Za-z]+[\s,]+\d{4}|\d{4}|present|current|till\s*date|till\s*now|ongoing)',
        re.IGNORECASE
    )

    # Collect non-overlapping intervals to avoid double-counting overlapping roles
    intervals = []
    for match in date_range_pattern.finditer(work_text):
        start = _parse_date(match.group(1))
        end = _parse_date(match.group(2))
        if start and end and end >= start:
            # Clamp end to today (future end dates are fine, e.g. "Present")
            end = min(end, today)
            if end > start:
                intervals.append((start, end))

    if not intervals:
        return None

    # Merge overlapping intervals to avoid double-counting parallel roles
    intervals.sort(key=lambda x: x[0])
    merged = [intervals[0]]
    for start, end in intervals[1:]:
        if start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))

    total_days = sum((e - s).days for s, e in merged)
    result = round(total_days / 365.25, 1)

    if 0 < result <= MAX_REASONABLE_YEARS:
        return result

    return None


# ── Chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str) -> List[str]:
    """Split text into overlapping chunks."""
    chunks, start = [], 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunks.append(text[start:end])
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


# ── CLI loader ────────────────────────────────────────────────────────────────

def load_cvs_from_folder(folder: str) -> List[CV]:
    """Load all PDF/DOCX files from a folder as CV objects."""
    folder_path = Path(folder)
    if not folder_path.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    supported = {".pdf", ".docx", ".doc"}
    files = [f for f in folder_path.iterdir() if f.suffix.lower() in supported]
    if not files:
        raise ValueError(f"No PDF/DOCX files found in: {folder}")

    cvs = []
    for file in sorted(files):
        text = _extract_text(str(file)).strip()
        if not text:
            continue
        chunks = _chunk_text(text)
        full_text = "\n\n---\n\n".join(chunks)
        candidate_name = file.stem.replace("_", " ").replace("-", " ").title()
        cvs.append(CV(
            id=str(uuid.uuid4()),
            candidate_name=candidate_name,
            raw_text=full_text,
        ))

    return cvs
